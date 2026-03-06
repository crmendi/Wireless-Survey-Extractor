import tkinter as tk
from tkinter import ttk, filedialog, messagebox, colorchooser
from docx import Document
from docx.shared import Inches
import zipfile
import json
import os
import csv
import re
import tempfile
import shutil
from PIL import Image, ImageDraw, ImageFont, ImageTk
import uuid
import matplotlib.pyplot as plt
import numpy as np
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import sys
import ctypes
from pypdf import PdfWriter

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

class WirelessSurveyExtractor:
    def __init__(self, root):
        self.root = root
        self.root.title("Wireless Survey Extractor")
        self.root.geometry("1100x750")
        self.root.state('zoomed')
        try:
            # Set App User Model ID for taskbar icon on Windows
            myappid = 'mycompany.myproduct.subproduct.version' # arbitrary string
            ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
        except Exception as e:
            print(f"Error setting AppUserModelID: {e}")

        try:
            ico_path = resource_path('icon.ico')
            png_path = resource_path('icon.png')
            if os.path.exists(ico_path):
                self.root.iconbitmap(ico_path)
            if os.path.exists(png_path):
                img = tk.PhotoImage(file=png_path)
                self.root.tk.call('wm', 'iconphoto', self.root._w, img)
        except Exception as e:
            print(f"Error setting icon: {e}")
        self.style = ttk.Style()
        self.setup_styles()
        self.create_icons()

        # Configuración inicial
        self.circle_radius = 10
        self.circle_color = "#E74C3C"
        self.text_color = "#2C3E50"
        self.font_size = 12
        self.outline_color = "#ECF0F1"
        self.outline_width = 2
        self.overwrite_all = False
        
        # Note settings
        self.note_color = "#000000"
        self.note_font_size = 12
        self.note_outline_color = "#FFFFFF"
        self.note_outline_width = 1
        
        self.ap_data = []
        self.aps_for_plotting = []
        self.notes_data = []
        self.note_counts = {}
        self.selected_files = []
        self.floor_order = []
        self.report_image_path = "Heimcore.png"
        self.current_language = tk.StringVar(value="es")
        self.translations = {
            "es": {
                "menu_file": "Archivo",
                "menu_export": "Exportar",
                "menu_tools": "Herramientas",
                "menu_help": "Ayuda",
                "menu_language": "Idioma",
                "select_esx": "Seleccionar Archivos ESX...",
                "load_project": "Cargar Proyecto...",
                "save_project": "Guardar Proyecto...",
                "exit": "Salir",
                "export_csv": "Exportar Datos a CSV...",
                "export_images": "Exportar Imágenes con APs...",
                "generate_word": "Generar Informe Word...",
                "merge_pdf": "Unir PDFs Existentes...",
                "settings": "Configuración...",
                "import_report_image": "Importar Imagen para Informe...",
                "tutorial": "Tutorial...",
                "about": "Acerca de...",
                "header_warning": "¡ATENCIÓN! Asegúrese de que el archivo .esx contenga ÚNICAMENTE diseños de APs Simulados (Levantamiento Predictivo). Los APs detectados en Site Surveys activos generarán fallos técnicos y visuales.",
                "file_filter": "Archivo:",
                "model_filter": "Modelo AP:",
                "floor_filter": "Piso:",
                "all": "Todos",
                "selected_files": "Archivos seleccionados: {files}",
                "tree_file": "Archivo",
                "tree_model": "Modelo",
                "tree_floor": "Piso",
                "tree_count": "Cantidad",
                "dialog_close": "Cerrar",
                "lang_es": "Español",
                "lang_en": "Inglés"
            },
            "en": {
                "menu_file": "File",
                "menu_export": "Export",
                "menu_tools": "Tools",
                "menu_help": "Help",
                "menu_language": "Language",
                "select_esx": "Select ESX Files...",
                "load_project": "Load Project...",
                "save_project": "Save Project...",
                "exit": "Exit",
                "export_csv": "Export Data to CSV...",
                "export_images": "Export Images with APs...",
                "generate_word": "Generate Word Report...",
                "merge_pdf": "Merge Existing PDFs...",
                "settings": "Settings...",
                "import_report_image": "Import Report Image...",
                "tutorial": "User Guide...",
                "about": "About...",
                "header_warning": "WARNING! Make sure the .esx file contains ONLY simulated AP designs (Predictive Survey). APs detected in active Site Surveys will generate technical and visual errors.",
                "file_filter": "File:",
                "model_filter": "AP Model:",
                "floor_filter": "Floor:",
                "all": "All",
                "selected_files": "Selected files: {files}",
                "tree_file": "File",
                "tree_model": "Model",
                "tree_floor": "Floor",
                "tree_count": "Count",
                "dialog_close": "Close",
                "lang_es": "Spanish",
                "lang_en": "English"
            }
        }
 
        self.create_menu()
        self.create_widgets()
        self.setup_bindings()
        self.refresh_language()

    def t(self, key, **kwargs):
        lang = self.current_language.get()
        text = self.translations.get(lang, {}).get(key, key)
        return text.format(**kwargs) if kwargs else text

    def is_english(self):
        return self.current_language.get() == "en"

    def tr(self, es, en):
        return en if self.is_english() else es

    def all_option(self):
        return self.t("all")

    def is_all_option(self, value):
        return value in [self.translations["es"]["all"], self.translations["en"]["all"]]

    def create_menu(self):
        self.menubar = tk.Menu(self.root, tearoff=0)
        self.root.config(menu=self.menubar)

        # File Menu
        self.file_menu = tk.Menu(self.menubar, tearoff=0)
        self.menubar.add_cascade(label=self.t("menu_file"), menu=self.file_menu)
        self.file_menu.add_command(label=self.t("select_esx"), command=self.load_esx, accelerator="Ctrl+O")
        self.file_menu.add_command(label=self.t("load_project"), command=self.load_project)
        self.file_menu.add_command(label=self.t("save_project"), command=self.save_project, accelerator="Ctrl+S")
        self.file_menu.add_separator()
        self.file_menu.add_command(label=self.t("exit"), command=self.root.quit)

        # Export Menu
        self.export_menu = tk.Menu(self.menubar, tearoff=0)
        self.menubar.add_cascade(label=self.t("menu_export"), menu=self.export_menu)
        self.export_menu.add_command(label=self.t("export_csv"), command=self.export_csv)
        self.export_menu.add_command(label=self.t("export_images"), command=self.export_images_with_aps)
        self.export_menu.add_command(label=self.t("generate_word"), command=self.generate_word_report)
        self.export_menu.add_command(label=self.t("merge_pdf"), command=self.merge_existing_pdfs)

        # Tools Menu
        self.tools_menu = tk.Menu(self.menubar, tearoff=0)
        self.menubar.add_cascade(label=self.t("menu_tools"), menu=self.tools_menu)
        self.tools_menu.add_command(label=self.t("settings"), command=self.show_settings_dialog)
        self.tools_menu.add_command(label=self.t("import_report_image"), command=self.import_report_image)

        # Help Menu
        self.help_menu = tk.Menu(self.menubar, tearoff=0)
        self.menubar.add_cascade(label=self.t("menu_help"), menu=self.help_menu)
        self.help_menu.add_command(label=self.t("tutorial"), command=self.show_tutorial_dialog)
        self.help_menu.add_command(label=self.t("about"), command=self.show_about_dialog)

        # Language Menu
        self.language_menu = tk.Menu(self.menubar, tearoff=0)
        self.menubar.add_cascade(label=self.t("menu_language"), menu=self.language_menu)
        self.language_menu.add_radiobutton(label=self.t("lang_es"), variable=self.current_language, value="es", command=self.refresh_language)
        self.language_menu.add_radiobutton(label=self.t("lang_en"), variable=self.current_language, value="en", command=self.refresh_language)


    def refresh_language(self):
        self.menubar.entryconfig(0, label=self.t("menu_file"))
        self.menubar.entryconfig(1, label=self.t("menu_export"))
        self.menubar.entryconfig(2, label=self.t("menu_tools"))
        self.menubar.entryconfig(3, label=self.t("menu_help"))
        self.menubar.entryconfig(4, label=self.t("menu_language"))

        self.file_menu.entryconfig(0, label=self.t("select_esx"))
        self.file_menu.entryconfig(1, label=self.t("load_project"))
        self.file_menu.entryconfig(2, label=self.t("save_project"))
        self.file_menu.entryconfig(4, label=self.t("exit"))

        self.export_menu.entryconfig(0, label=self.t("export_csv"))
        self.export_menu.entryconfig(1, label=self.t("export_images"))
        self.export_menu.entryconfig(2, label=self.t("generate_word"))
        self.export_menu.entryconfig(3, label=self.t("merge_pdf"))

        self.tools_menu.entryconfig(0, label=self.t("settings"))
        self.tools_menu.entryconfig(1, label=self.t("import_report_image"))

        self.help_menu.entryconfig(0, label=self.t("tutorial"))
        self.help_menu.entryconfig(1, label=self.t("about"))

        self.language_menu.entryconfig(0, label=self.t("lang_es"))
        self.language_menu.entryconfig(1, label=self.t("lang_en"))

        if hasattr(self, "warning_label"):
            self.warning_label.config(text=self.t("header_warning"))
        if hasattr(self, "filter_labels"):
            self.filter_labels['archivo'].config(text=self.t("file_filter"))
            self.filter_labels['modelo'].config(text=self.t("model_filter"))
            self.filter_labels['piso'].config(text=self.t("floor_filter"))

        self.tree.heading('Archivo', text=self.t("tree_file"))
        self.tree.heading('Modelo', text=self.t("tree_model"))
        self.tree.heading('Piso', text=self.t("tree_floor"))
        self.tree.heading('Cantidad', text=self.t("tree_count"))

        self.update_combobox(self.combo_archivo, {entry[0] for entry in self.ap_data}, self.combo_archivo.get())
        self.update_combobox(self.combo_modelo, {entry[1] for entry in self.ap_data}, self.combo_modelo.get())
        self.update_combobox(self.combo_piso, {entry[2] for entry in self.ap_data}, self.combo_piso.get())

        if self.selected_files:
            self.selected_files_label.config(text=self.t("selected_files", files=', '.join(os.path.basename(f) for f in self.selected_files)))

    def show_tutorial_dialog(self):
        tutorial_win = tk.Toplevel(self.root)
        tutorial_win.title(self.tr("Tutorial de Wireless Survey Extractor", "Wireless Survey Extractor Tutorial"))
        tutorial_win.geometry("800x650")
        tutorial_win.transient(self.root)
        tutorial_win.grab_set()

        main_frame = ttk.Frame(tutorial_win, padding=15)
        main_frame.pack(fill=tk.BOTH, expand=True)

        text_frame = ttk.Frame(main_frame)
        text_frame.pack(fill=tk.BOTH, expand=True)

        text_widget = tk.Text(text_frame, wrap=tk.WORD, font=("Segoe UI", 10), relief=tk.FLAT, padx=10, pady=10)
        scrollbar = ttk.Scrollbar(text_frame, orient="vertical", command=text_widget.yview)
        text_widget.configure(yscrollcommand=scrollbar.set)

        text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Define tag for bold text
        bold_font = ("Segoe UI", 10, "bold")
        text_widget.tag_configure("bold", font=bold_font)        tutorial_text = self.tr(
            """
**Bienvenido al Tutorial de Wireless Survey Extractor**

Esta guia te mostrara como utilizar todas las funcionalidades de la aplicacion.

**1. Cargar Archivos (Solo Disenos Predictivos)**

-   **Seleccionar Archivos ESX... (Ctrl+O):** Ve a `Archivo > Seleccionar Archivos ESX...`.
-   **AVISO:** La aplicacion procesara SOLO los APs simulados (diseno predictivo). Los APs detectados en Site Surveys activos pueden generar errores visuales o lecturas erraticas.
-   **Cargar Proyecto...:** Usa `Archivo > Cargar Proyecto...` para abrir un archivo `.aproj` guardado previamente.

**2. Guardar un Proyecto**

-   **Guardar Proyecto... (Ctrl+S):** Ve a `Archivo > Guardar Proyecto...` para guardar tu sesion actual.

**3. Filtrar Datos**

-   Puedes filtrar por `Archivo`, `Modelo de AP` o `Piso` usando los desplegables.

**4. Exportar Resultados**

El menu `Exportar` contiene las opciones para generar entregables:

-   **Exportar Datos a CSV...**
-   **Exportar Imagenes con APs...**
-   **Generar Informe Word...**

**5. Configuracion y Personalizacion**

-   **Configuracion...:** Ajusta colores, tamano de fuente y radio de circulos.
-   **Importar Imagen para Informe...:** Selecciona una imagen para portada.

**6. Ayuda**

-   **Acerca de...:** Muestra informacion de la aplicacion.
""",
            """
**Welcome to the Wireless Survey Extractor Tutorial**

This guide shows how to use the main features of the app.

**1. Load Files (Predictive Designs Only)**

-   **Select ESX Files... (Ctrl+O):** Go to `File > Select ESX Files...`.
-   **NOTICE:** The app should process ONLY simulated AP designs (predictive survey). APs detected in active Site Surveys can cause visual errors or unstable readings.
-   **Load Project...:** Use `File > Load Project...` to open a previously saved `.aproj` project.

**2. Save a Project**

-   **Save Project... (Ctrl+S):** Go to `File > Save Project...` to save your current session.

**3. Filter Data**

-   You can filter by `File`, `AP Model`, or `Floor` with the dropdowns.

**4. Export Results**

The `Export` menu contains output options:

-   **Export Data to CSV...**
-   **Export Images with APs...**
-   **Generate Word Report...**

**5. Settings and Customization**

-   **Settings...:** Adjust colors, font size, and circle radius.
-   **Import Report Image...:** Select an image for the report cover.

**6. Help**

-   **About...:** Shows application information.
"""
        )
        # Process and insert text with bold tags
        for line in tutorial_text.strip().split('\n'):
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    text_widget.insert(tk.END, part[2:-2], "bold")
                else:
                    text_widget.insert(tk.END, part)
            text_widget.insert(tk.END, '\n')

        text_widget.config(state=tk.DISABLED) # Make text read-only

        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(pady=(10, 0))
        ttk.Button(btn_frame, text=self.tr("Cerrar", "Close"), command=tutorial_win.destroy).pack()

    def import_report_image(self):
        path = filedialog.askopenfilename(
            title=self.tr("Seleccionar imagen para el informe", "Select report image"),
            filetypes=(("PNG files", "*.png"), ("JPEG files", "*.jpg;*.jpeg"), (self.tr("Todos los archivos", "All files"), "*.*"))
        )
        if path:
            self.report_image_path = path
            messagebox.showinfo(self.tr("Exito", "Success"), self.tr(f"Imagen '{os.path.basename(path)}' seleccionada para los informes.", f"Image '{os.path.basename(path)}' selected for reports."))
    def show_about_dialog(self):
        license_text = self.tr(
            """
Wireless Survey Extractor - Version 1.0
Creado por: Christian Mendivelso

----------------------------------
Licencia Apache 2.0 (Codigo Abierto)
----------------------------------

Copyright (c) 2025 Christian Mendivelso

Este software es de codigo abierto. Tienes total libertad para:
- Usarlo con cualquier proposito (incluyendo comercial).
- Modificarlo y adaptarlo a tus necesidades.
- Distribuirlo.

BAJO LA SIGUIENTE CONDICION OBLIGATORIA:
Debes dar el credito correspondiente al autor original (Christian Mendivelso) e incluir una copia o enlace a la licencia original en cualquier distribucion, modificacion o trabajo derivado.

EL SOFTWARE SE PROPORCIONA "TAL CUAL", SIN GARANTIAS DE NINGUN TIPO.
""",
            """
Wireless Survey Extractor - Version 1.0
Created by: Christian Mendivelso

----------------------------------
Apache 2.0 License (Open Source)
----------------------------------

Copyright (c) 2025 Christian Mendivelso

This software is open source. You are free to:
- Use it for any purpose (including commercial use).
- Modify and adapt it to your needs.
- Distribute it.

UNDER THE FOLLOWING MANDATORY CONDITION:
You must provide proper credit to the original author (Christian Mendivelso) and include a copy of or link to the original license in any distribution, modification, or derivative work.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTIES OF ANY KIND.
"""
        )
        messagebox.showinfo(
            self.tr("Acerca de Wireless Survey Extractor", "About Wireless Survey Extractor"),
            license_text
        )

    def setup_styles(self):
        self.style.theme_use('clam')
        self.style.configure('.', background='#ECF0F1', foreground='#2C3E50')
        self.style.configure('TButton', font=('Segoe UI', 10), padding=8,
                           borderwidth=0, focuscolor='#BDC3C7')
        self.style.map('TButton',
                     background=[('active', '#BDC3C7'), ('!disabled', '#ECF0F1')],
                     foreground=[('!disabled', '#2C3E50')])
        self.style.configure('Header.TLabel', font=('Segoe UI', 14, 'bold'),
                           foreground='#2C3E50', background='#ECF0F1')
        self.style.configure('Warning.TLabel', font=('Segoe UI', 11, 'bold'),
                           foreground='red', background='#ECF0F1')
        self.style.configure('Treeview', font=('Segoe UI', 10), rowheight=28,
                           fieldbackground='#FFFFFF')
        self.style.configure('Treeview.Heading', font=('Segoe UI', 10, 'bold'),
                           background='#3498DB', foreground='#FFFFFF')
        self.style.map('Treeview',
                     background=[('selected', '#2980B9')],
                     foreground=[('selected', 'white')])

    def create_icons(self):
        icon_config = {
            'open': ('📂', '#27AE60'),
            'settings': ('⚙', '#34495E'),
            'save': ('💾', '#2980B9'),
            'load': ('📁', '#F39C12'),
            'export': ('📤', '#2ECC71'),
            'report': ('📄', '#9B59B6'),
            'color': ('🎨', '#E74C3C')
        }
        self.icons = {k: self.generate_icon(v[0], v[1]) for k, v in icon_config.items()}

    def generate_icon(self, symbol, bg_color, size=(32, 32)):
        img = Image.new('RGBA', size, (0,0,0,0))
        draw = ImageDraw.Draw(img)
        draw.rounded_rectangle([0, 0, size[0]-1, size[1]-1], radius=6, fill=bg_color)
        try:
            font = ImageFont.truetype("seguiemj.ttf", 18)
        except:
            font = ImageFont.load_default()
        
        # Usar textbbox en lugar de textsize
        bbox = draw.textbbox((0, 0), symbol, font=font)
        w = bbox[2] - bbox[0]  # Ancho del texto
        h = bbox[3] - bbox[1]  # Alto del texto
        
        draw.text(((size[0]-w)/2, (size[1]-h)/2-2), 
                symbol, font=font, fill='white')
        return ImageTk.PhotoImage(img)

    def center_window(self):
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f'+{x}+{y}')

    def create_widgets(self):
        main_frame = ttk.Frame(self.root, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Header
        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 10))
        ttk.Label(header_frame, text="Wireless Survey Extractor", style='Header.TLabel').pack(side=tk.LEFT)

        # Warning Message
        self.warning_label = ttk.Label(
            main_frame,
            text=self.t("header_warning"),
            style="Warning.TLabel",
            wraplength=1000
        )
        self.warning_label.pack(fill=tk.X, pady=(0, 10))

        # Panel principal
        content_frame = ttk.Frame(main_frame)
        content_frame.pack(fill=tk.BOTH, expand=True)

        # Filtros
        filter_frame = ttk.Frame(content_frame)
        filter_frame.pack(fill=tk.X, pady=10)
        
        self.filter_labels = {}
        filters = [
            ("file_filter", 'archivo', 0),
            ("model_filter", 'modelo', 2),
            ("floor_filter", 'piso', 4)
        ]
        
        for label_key, combo_name, col in filters:
            lbl = ttk.Label(filter_frame, text=self.t(label_key))
            lbl.grid(row=0, column=col, padx=5, sticky=tk.W)
            self.filter_labels[combo_name] = lbl
            combo = ttk.Combobox(filter_frame, state="readonly", width=22)
            combo.grid(row=0, column=col+1, padx=5, sticky=tk.EW)
            combo.set(self.all_option())
            setattr(self, f'combo_{combo_name}', combo)

        # Tabla de resultados
        tree_frame = ttk.Frame(content_frame)
        tree_frame.pack(fill=tk.BOTH, expand=True)
        
        self.tree = ttk.Treeview(tree_frame, columns=('Archivo', 'Modelo', 'Piso', 'Cantidad'),
                               show='headings', selectmode='browse')
        
        columns = [
            ('Archivo', 250),
            ('Modelo', 300),
            ('Piso', 150),
            ('Cantidad', 100)
        ]
        
        for col, width in columns:
            self.tree.heading(col, text=col)
            self.tree.column(col, width=width, anchor=tk.W if col != 'Cantidad' else tk.CENTER)
        
        vsb = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=self.tree.yview)
        hsb = ttk.Scrollbar(tree_frame, orient=tk.HORIZONTAL, command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        
        self.tree.grid(row=0, column=0, sticky=tk.NSEW)
        vsb.grid(row=0, column=1, sticky=tk.NS)
        hsb.grid(row=1, column=0, sticky=tk.EW)
        tree_frame.grid_rowconfigure(0, weight=1)
        tree_frame.grid_columnconfigure(0, weight=1)

        # Status Bar
        self.status_bar = ttk.Label(main_frame, text="2025 - Christian Mendivelso", anchor=tk.W)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)

        self.selected_files_label = ttk.Label(main_frame, wraplength=1000)
        self.selected_files_label.pack(pady=10)

    def setup_bindings(self):
        self.combo_archivo.bind('<<ComboboxSelected>>', self.apply_filters)
        self.combo_modelo.bind('<<ComboboxSelected>>', self.apply_filters)
        self.combo_piso.bind('<<ComboboxSelected>>', self.apply_filters)

    def show_settings_dialog(self):
        settings_win = tk.Toplevel(self.root)
        settings_win.title(self.tr("Configuracion de Visualizacion", "Display Settings"))
        settings_win.geometry("800x600")
        self.center_window()
        
        main_frame = ttk.Frame(settings_win, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Panel de controles
        control_frame = ttk.Frame(main_frame)
        control_frame.pack(side=tk.LEFT, fill=tk.Y, padx=10)
                settings = [
            (self.tr("Radio del circulo:", "Circle radius:"), 'circle_radius', self.circle_radius),
            (self.tr("Tamano de fuente:", "Font size:"), 'font_size', self.font_size),
            (self.tr("Grosor delineado:", "Outline width:"), 'outline_width', self.outline_width),
            (self.tr("Tamano de fuente de nota:", "Note font size:"), 'note_font_size', self.note_font_size),
            (self.tr("Grosor delineado de nota:", "Note outline width:"), 'note_outline_width', self.note_outline_width)
        ]
        
        self.entries = {}
        for i, (label, name, value) in enumerate(settings):
            ttk.Label(control_frame, text=label).grid(row=i, column=0, pady=5, sticky=tk.W)
            entry = ttk.Entry(control_frame, width=8)
            entry.insert(0, str(value))
            entry.grid(row=i, column=1, pady=5, sticky=tk.EW)
            self.entries[name] = entry
                color_buttons = [
            (self.tr("Color del circulo", "Circle color"), 'circle_color', self.circle_color),
            (self.tr("Color del texto", "Text color"), 'text_color', self.text_color),
            (self.tr("Color delineado", "Outline color"), 'outline_color', self.outline_color),
            (self.tr("Color de la nota", "Note color"), 'note_color', self.note_color),
            (self.tr("Color delineado de nota", "Note outline color"), 'note_outline_color', self.note_outline_color)
        ]
        
        for i, (text, color_type, color) in enumerate(color_buttons, len(settings)):
            btn = ttk.Button(control_frame, text=text, image=self.icons['color'],
                           compound=tk.LEFT, command=lambda ct=color_type: self.choose_color(ct))
            btn.grid(row=i, columnspan=2, pady=5, sticky=tk.W)
        
        # Panel de vista previa
        preview_frame = ttk.Frame(main_frame)
        preview_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        
        ttk.Label(preview_frame, text=self.tr("Vista Previa", "Preview")).pack(pady=5)
        self.preview_canvas = tk.Canvas(preview_frame, bg='white', bd=2, relief=tk.GROOVE)
        self.preview_canvas.pack(fill=tk.BOTH, expand=True)
        
        # Botones de acción
        btn_frame = ttk.Frame(control_frame)
        btn_frame.grid(row=6, columnspan=2, pady=20)
        ttk.Button(btn_frame, text=self.tr("Aplicar", "Apply"), command=self.save_settings).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text=self.tr("Cerrar", "Close"), command=settings_win.destroy).pack(side=tk.LEFT, padx=5)
        
        self.generate_preview()

    def generate_preview(self):
        self.preview_canvas.delete("all")
        try:
            radius = int(self.entries['circle_radius'].get())
            font_size = int(self.entries['font_size'].get())
            outline_width = int(self.entries['outline_width'].get())
            note_font_size = int(self.entries['note_font_size'].get())
            note_outline_width = int(self.entries['note_outline_width'].get())
        except ValueError:
            return
        
        preview_img = Image.new('RGB', (400, 300), color='white')
        draw = ImageDraw.Draw(preview_img)
        
        sample_aps = [
            {'x': 100, 'y': 100, 'name': 'AP-1'},
            {'x': 300, 'y': 200, 'name': 'AP-2'}
        ]
        
        for ap in sample_aps:
            x = ap['x']
            y = ap['y']
            
            # Dibujar círculo
            draw.ellipse((x - radius, y - radius, x + radius, y + radius),
                        fill=self.circle_color, outline="black")
            
            # Configurar fuente
            try:
                font = ImageFont.truetype("arialbd.ttf", font_size)
            except IOError:
                font = ImageFont.load_default()
            
            # Dibujar texto con borde
            text_position = (x + radius + 2, y - radius)
            if outline_width > 0:
                for dx in range(-outline_width, outline_width + 1):
                    for dy in range(-outline_width, outline_width + 1):
                        if dx != 0 or dy != 0:
                            draw.text(
                                (text_position[0] + dx, text_position[1] + dy),
                                ap['name'],
                                fill=self.outline_color,
                                font=font
                            )
            
            # Dibujar texto principal
            draw.text(text_position, ap['name'], 
                    fill=self.text_color, font=font)
        
        # Preview for notes
        note_font = ImageFont.truetype("arial.ttf", note_font_size)
        if note_outline_width > 0:
            for dx in range(-note_outline_width, note_outline_width + 1):
                for dy in range(-note_outline_width, note_outline_width + 1):
                    if dx != 0 or dy != 0:
                        draw.text(
                            (50 + dx, 250 + dy),
                            "Ejemplo de nota",
                            fill=self.note_outline_color,
                            font=note_font
                        )
        note_example = self.tr("Ejemplo de nota", "Sample note")
        draw.text((50, 250), note_example, fill=self.note_color, font=note_font)

        self.tk_preview = ImageTk.PhotoImage(preview_img)
        self.preview_canvas.create_image(0, 0, anchor=tk.NW, image=self.tk_preview)

    def save_settings(self):
        try:
            self.circle_radius = int(self.entries['circle_radius'].get())
            self.font_size = int(self.entries['font_size'].get())
            self.outline_width = int(self.entries['outline_width'].get())
            self.note_font_size = int(self.entries['note_font_size'].get())
            self.note_outline_width = int(self.entries['note_outline_width'].get())
            self.generate_preview()
        except ValueError:
            messagebox.showerror("Error", "Valores numéricos inválidos")

    def choose_color(self, color_type):
        color = colorchooser.askcolor()[1]
        if color:
            if color_type == 'circle_color':
                self.circle_color = color
            elif color_type == 'text_color':
                self.text_color = color
            elif color_type == 'outline_color':
                self.outline_color = color
            elif color_type == 'note_color':
                self.note_color = color
            elif color_type == 'note_outline_color':
                self.note_outline_color = color
            self.generate_preview()

    def load_esx(self):
        file_paths = filedialog.askopenfilenames(
            title=self.tr('Seleccionar archivo(s) .esx', 'Select .esx file(s)'),
            filetypes=(("Extractor files", ".esx"), ("Todos los archivos", ".*"))
        )
        
        if file_paths:
            self.ap_data = []
            self.aps_for_plotting = []
            self.notes_data = []
            self.note_counts = {}
            self.selected_files = file_paths
            self.selected_files_label.config(
                text=self.t("selected_files", files=', '.join(os.path.basename(f) for f in file_paths))
            )
            
            for file_path in file_paths:
                try:
                    self.process_esx_file(file_path)
                except Exception as e:
                    messagebox.showerror("Error", f"Error procesando {os.path.basename(file_path)}:\n{str(e)}")
            
            self.combo_archivo.set(self.all_option())
            self.combo_modelo.set(self.all_option())
            self.combo_piso.set(self.all_option())
            
            self.ap_data.sort(key=lambda x: self.floor_and_block_sort_key((x[0], x[2])))
            self.apply_filters(force_update=True)

    def process_esx_file(self, file_path):
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            filename = os.path.basename(file_path)
            
            floor_data = self.load_json_from_zip(zip_ref, 'floorPlans.json').get('floorPlans', [])
            floor_map = {}
            for floor in floor_data:
                try:
                    floor_id = floor['id']
                    floor_map[floor_id] = {
                        'number': self.extract_floor_number(floor.get('name', '')),
                        'width': float(floor.get('width', 1.0)),
                        'height': float(floor.get('height', 1.0)),
                        'image_id': floor.get('imageId'),
                    }
                except Exception as e:
                    print(f"Error procesando piso: {str(e)}")
                    continue

            ap_data = self.load_json_from_zip(zip_ref, 'accessPoints.json').get('accessPoints', [])
            ap_counts = {}
            
            for ap in ap_data:
                try:
                    location = ap.get('location', {})
                    coord = location.get('coord', {})
                    floor_plan_id = location.get('floorPlanId')
                    floor_info = floor_map.get(floor_plan_id)
                    
                    if not floor_info or not coord:
                        continue
                    
                    model = ap.get('model', 'Desconocido')
                    floor_number = floor_info['number']
                    
                    key = (model, floor_number)
                    ap_counts[key] = ap_counts.get(key, 0) + 1
                    
                    self.aps_for_plotting.append({
                        'filename': filename,
                        'model': model,
                        'floor': floor_number,
                        'x': float(coord.get('x', 0)),
                        'y': float(coord.get('y', 0)),
                        'name': ap.get('name', 'AP')
                    })
                        
                except Exception as e:
                    print(f"Error procesando AP: {str(e)}")
                    continue
            
            for (model, floor), count in ap_counts.items():
                self.ap_data.append((filename, model, floor, count))

            # Extract notes
            notes_json = self.load_json_from_zip(zip_ref, 'notes.json').get('notes', [])
            picture_notes_json = self.load_json_from_zip(zip_ref, 'pictureNotes.json').get('pictureNotes', [])

            notes_map = {note['id']: note['text'] for note in notes_json}

            self.note_counts[filename] = {}
            for pic_note in picture_notes_json:
                try:
                    location = pic_note.get('location', {})
                    coord = location.get('coord', {})
                    floor_plan_id = location.get('floorPlanId')
                    floor_info = floor_map.get(floor_plan_id)

                    if not floor_info or not coord:
                        continue

                    for note_id in pic_note.get('noteIds', []):
                        if note_id in notes_map:
                            note_text = notes_map[note_id]
                            
                            # Reemplazar "reubicacion" por "traslado"
                            note_text = re.sub(r'reubicacion', 'traslado', note_text, flags=re.IGNORECASE)

                            self.notes_data.append({
                                'id': str(uuid.uuid4()),
                                'filename': filename,
                                'floor': floor_info['number'],
                                'x': float(coord.get('x', 0)),
                                'y': float(coord.get('y', 0)),
                                'text': note_text
                            })
                            
                            note_type = note_text.lower()
                            if note_type in ["existente", "nuevo", "traslado"]:
                                floor_number = floor_info['number']
                                if floor_number not in self.note_counts[filename]:
                                    self.note_counts[filename][floor_number] = {"existente": 0, "nuevo": 0, "traslado": 0}
                                self.note_counts[filename][floor_number][note_type] += 1

                except Exception as e:
                    print(f"Error procesando nota: {str(e)}")
                    continue

    def export_images_with_aps(self):
        if not self.aps_for_plotting:
            messagebox.showwarning("Advertencia", "No hay APs para graficar")
            return
        
        save_dir = filedialog.askdirectory(title="Seleccionar carpeta para guardar las imágenes")
        if not save_dir:
            return
        
        try:
            self.overwrite_all = False
            for file_path in self.selected_files:
                filename = os.path.basename(file_path)
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    images_data = self.load_json_from_zip(zip_ref, 'images.json').get('images', [])
                    image_map = {img['id']: img for img in images_data}
                    
                    floor_plans = self.load_json_from_zip(zip_ref, 'floorPlans.json').get('floorPlans', [])
                    
                    for floor in floor_plans:
                        self.process_floor_image(zip_ref, floor, image_map, filename, save_dir)
            
            messagebox.showinfo("Éxito", "Imágenes exportadas correctamente")
        except Exception as e:
            messagebox.showerror("Error", f"Error al exportar imágenes:\n{str(e)}")
        finally:
            self.overwrite_all = False

    def process_floor_image(self, zip_ref, floor, image_map, filename, save_dir):
        image_id = floor.get('imageId')
        if not image_id:
            return
        
        try:
            with zip_ref.open(f"image-{image_id}") as img_file:
                image = Image.open(img_file).convert("RGB")
            
            image_info = image_map.get(image_id, {})
            original_width = image_info.get('resolutionWidth', image.width)
            original_height = image_info.get('resolutionHeight', image.height)
            
            floor_name = floor.get('name', '')
            floor_number = self.extract_floor_number(floor_name)
            
            draw = ImageDraw.Draw(image)
            
            safe_name = re.sub(r'[\\/*?:"<>|]', '_', floor_name)
            output_path = os.path.join(save_dir, f"{filename}_{safe_name}_APs.png")
            
            if os.path.exists(output_path):
                if not self.overwrite_all:
                    response = self.ask_overwrite(os.path.basename(output_path))
                    if response == "yes_all":
                        self.overwrite_all = True
                    elif response == "yes":
                        pass
                    else:
                        return
            
            ap_count = 0
            for ap in self.aps_for_plotting:
                if ap['filename'] != filename or ap['floor'] != floor_number:
                    continue
                
                try:
                    x_final = ap['x']
                    y_final = ap['y']
                    
                    draw.ellipse((x_final-self.circle_radius, 
                                y_final-self.circle_radius,
                                x_final+self.circle_radius, 
                                y_final+self.circle_radius), 
                               fill=self.circle_color, 
                               outline="black")
                    
                    try:
                        font = ImageFont.truetype("arialbd.ttf", self.font_size)
                    except IOError:
                        font = ImageFont.load_default()
                    
                    text_position = (x_final+self.circle_radius+2, 
                                   y_final-self.circle_radius)
                    
                    if self.outline_width > 0:
                        for dx in range(-self.outline_width, self.outline_width+1):
                            for dy in range(-self.outline_width, self.outline_width+1):
                                if dx != 0 or dy != 0:
                                    draw.text(
                                        (text_position[0] + dx, 
                                         text_position[1] + dy),
                                        ap['name'],
                                        fill=self.outline_color,
                                        font=font
                                    )
                    
                    draw.text(text_position, 
                            ap['name'], 
                            fill=self.text_color, 
                            font=font)
                    
                    ap_count += 1
                
                except Exception as e:
                    print(f"Error dibujando AP {ap.get('name', '')}: {str(e)}")
                    continue
            
            # Get AP label bounding boxes to avoid overlap
            ap_bboxes = []
            font = ImageFont.truetype("arialbd.ttf", self.font_size)
            for ap in self.aps_for_plotting:
                 if ap['filename'] == filename and ap['floor'] == floor_number:
                    text_position = (ap['x'] + self.circle_radius + 2, ap['y'] - self.circle_radius)
                    bbox = draw.textbbox(text_position, ap['name'], font=font)
                    ap_bboxes.append(bbox)

            floor_notes = [n for n in self.notes_data if n['filename'] == filename and n['floor'] == floor_number]
            
            for note in floor_notes:
                note_font = ImageFont.truetype("arial.ttf", self.note_font_size)
                note_bbox = draw.textbbox((note['x'], note['y']), note['text'], font=note_font)

                # Collision detection and avoidance
                attempts = 0
                while any(self.bboxes_overlap(note_bbox, ap_bbox) for ap_bbox in ap_bboxes) and attempts < 20:
                    # Move note down and to the right to find a clear spot
                    note['x'] += 5
                    note['y'] += 5
                    note_bbox = draw.textbbox((note['x'], note['y']), note['text'], font=note_font)
                    attempts += 1
                
                self.draw_note(draw, note)
            
            image.save(output_path)
        
        except Exception as e:
            print(f"Error procesando piso {floor_name}: {str(e)}")

    def draw_note(self, draw, note):
        x = note['x']
        y = note['y']
        text = note['text']
        
        try:
            font = ImageFont.truetype("arial.ttf", self.note_font_size)
        except IOError:
            font = ImageFont.load_default()
        
        # Simple note representation: text
        if self.note_outline_width > 0:
            for dx in range(-self.note_outline_width, self.note_outline_width + 1):
                for dy in range(-self.note_outline_width, self.note_outline_width + 1):
                    if dx != 0 or dy != 0:
                        draw.text(
                            (x + dx, y + dy),
                            text,
                            fill=self.note_outline_color,
                            font=font
                        )
        draw.text((x, y), text, fill=self.note_color, font=font)

    def ask_overwrite(self, filename):
        dialog = tk.Toplevel(self.root)
        dialog.title("Archivo existente")
        dialog.transient(self.root)
        dialog.grab_set()
        
        msg = f"El archivo {filename} ya existe.\n¿Qué deseas hacer?"
        lbl = ttk.Label(dialog, text=msg)
        lbl.pack(pady=10, padx=20)
        
        response = None
        
        def set_response(res):
            nonlocal response
            response = res
            dialog.destroy()
        
        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(pady=10)
        
        ttk.Button(btn_frame, text="Sobreescribir", 
                 command=lambda: set_response("yes")).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Sobreescribir todos", 
                 command=lambda: set_response("yes_all")).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Cancelar", 
                 command=lambda: set_response("no")).pack(side=tk.LEFT, padx=5)
        
        self.root.wait_window(dialog)
        return response

    def load_json_from_zip(self, zip_ref, filename):
        try:
            with zip_ref.open(filename) as f:
                return json.load(f)
        except:
            return {}

    def extract_floor_number(self, floor_name):
        lower_name = floor_name.lower()
        if 'sotano' in lower_name:
            return 'Sotano'
        if 'mezanine' in lower_name or 'mezzanine' in lower_name:
            return 'Mezanine'
        match = re.search(r'\d+', floor_name)
        return match.group() if match else floor_name if floor_name else 'Unknown'

    def floor_sort_key(self, floor_str):
        if isinstance(floor_str, str):
            lower_str = floor_str.lower()
            if 'sotano' in lower_str:
                return -1
            if 'mezanine' in lower_str or 'mezzanine' in lower_str:
                return 0.5
        try:
            return int(floor_str)
        except (ValueError, TypeError):
            return float('inf')

    def apply_filters(self, event=None, force_update=False):
        current_archivo = self.combo_archivo.get()
        current_modelo = self.combo_modelo.get()
        current_piso = self.combo_piso.get()
        
        filtered = [
            entry for entry in self.ap_data
            if ((self.is_all_option(current_archivo) or current_archivo == entry[0])) and
               ((self.is_all_option(current_modelo) or current_modelo == entry[1])) and
               ((self.is_all_option(current_piso) or current_piso == entry[2]))
        ]
        self.update_table(filtered)
        self.update_combos(current_archivo, current_modelo, current_piso, force=force_update)

    def update_combos(self, current_archivo, current_modelo, current_piso, force=False):
        files = {entry[0] for entry in self.ap_data}
        models = {entry[1] for entry in self.ap_data}
        floors = {entry[2] for entry in self.ap_data}
        
        self.update_combobox(self.combo_archivo, files, current_archivo)
        self.update_combobox(self.combo_modelo, models, current_modelo)
        self.update_combobox(self.combo_piso, floors, current_piso)
        
        if force:
            self.combo_archivo.event_generate("<<ComboboxSelected>>")
            self.combo_modelo.event_generate("<<ComboboxSelected>>")
            self.combo_piso.event_generate("<<ComboboxSelected>>")

    def update_combobox(self, combo, options, current):
        options = sorted(options, key=lambda x: (self.floor_sort_key(x) if combo == self.combo_piso else x))
        combo['values'] = [self.all_option()] + list(options)
        combo.set(current if current in options or self.is_all_option(current) else self.all_option())

    def update_table(self, data):
        self.tree.delete(*self.tree.get_children())
        for entry in data:
            self.tree.insert('', tk.END, values=entry)

    def export_csv(self):
        if not self.ap_data:
            messagebox.showwarning(self.tr('Advertencia', 'Warning'), self.tr('No hay datos para exportar', 'No data to export'))
            return
        
        save_path = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=(("CSV files", ".csv"), ("Todos los archivos", ".*"))
        )
        
        if save_path:
            try:
                with open(save_path, 'w', newline='', encoding='utf-8') as f:
                    writer = csv.writer(f)
                    writer.writerow(['Archivo', 'Modelo', 'Piso', 'Cantidad'])
                    writer.writerows(self.ap_data)
                messagebox.showinfo("Éxito", "CSV exportado correctamente")
            except Exception as e:
                messagebox.showerror("Error", f"Error al exportar CSV:\n{str(e)}")

    def save_project(self):
        if not self.selected_files:
            messagebox.showwarning("Advertencia", "No hay proyectos para guardar")
            return
        
        save_path = filedialog.asksaveasfilename(
            defaultextension=".aproj",
            filetypes=(("Archivos de proyecto", "*.aproj"), ("Todos los archivos", "*.*"))
        )
        
        if save_path:
            try:
                project_data = {
                    "version": 1,
                    "selected_files": self.selected_files,
                    "config": {
                        "circle_radius": self.circle_radius,
                        "circle_color": self.circle_color,
                        "text_color": self.text_color,
                        "font_size": self.font_size,
                        "outline_color": self.outline_color,
                        "outline_width": self.outline_width,
                        "note_color": self.note_color,
                        "note_font_size": self.note_font_size,
                        "note_outline_color": self.note_outline_color,
                        "note_outline_width": self.note_outline_width
                    },
                    "filters": {
                        "archivo": self.combo_archivo.get(),
                        "modelo": self.combo_modelo.get(),
                        "piso": self.combo_piso.get()
                    }
                }
                
                with open(save_path, 'w') as f:
                    json.dump(project_data, f, indent=4)
                
                messagebox.showinfo("Éxito", "Proyecto guardado correctamente")
            except Exception as e:
                messagebox.showerror("Error", f"Error al guardar proyecto:\n{str(e)}")

    def load_project(self):
        file_path = filedialog.askopenfilename(
            filetypes=(("Archivos de proyecto", "*.aproj"), ("Todos los archivos", "*.*"))
        )
        
        if file_path:
            try:
                with open(file_path, 'r') as f:
                    project_data = json.load(f)
                
                if project_data.get("version") != 1:
                    messagebox.showerror("Error", "Versión de proyecto no soportada")
                    return
                
                self.selected_files = project_data["selected_files"]
                self.selected_files_label.config(
                    text=self.t("selected_files", files=', '.join(os.path.basename(f) for f in self.selected_files))
                )
                
                self.ap_data = []
                self.aps_for_plotting = []
                self.notes_data = []
                self.note_counts = {}
                
                missing_files = []
                for file_path in self.selected_files:
                    if not os.path.exists(file_path):
                        missing_files.append(os.path.basename(file_path))
                        continue
                    try:
                        self.process_esx_file(file_path)
                    except Exception as e:
                        messagebox.showerror("Error", f"Error procesando {os.path.basename(file_path)}:\n{str(e)}")
                
                if missing_files:
                    messagebox.showwarning("Archivos faltantes", 
                                        f"No se encontraron: {', '.join(missing_files)}")
                
                config = project_data.get("config", {})
                self.circle_radius = config.get("circle_radius", 8)
                self.circle_color = config.get("circle_color", "red")
                self.text_color = config.get("text_color", "black")
                self.font_size = config.get("font_size", 14)
                self.outline_color = config.get("outline_color", "white")
                self.outline_width = config.get("outline_width", 1)
                self.note_color = config.get("note_color", "#000000")
                self.note_font_size = config.get("note_font_size", 12)
                self.note_outline_color = config.get("note_outline_color", "#FFFFFF")
                self.note_outline_width = config.get("note_outline_width", 1)
                
                filters = project_data.get("filters", {})
                self.combo_archivo.set(filters.get("archivo", self.all_option()))
                self.combo_modelo.set(filters.get("modelo", self.all_option()))
                self.combo_piso.set(filters.get("piso", self.all_option()))
                
                self.ap_data.sort(key=lambda x: self.floor_and_block_sort_key((x[0], x[2])))
                self.apply_filters(force_update=True)
                
                messagebox.showinfo("Éxito", "Proyecto cargado correctamente")
            
            except Exception as e:
                messagebox.showerror("Error", f"Error al cargar proyecto:\n{str(e)}")

    def generate_word_report(self):
        if not self.aps_for_plotting:
            messagebox.showwarning(self.tr('Advertencia', 'Warning'), self.tr('No hay APs para generar informe', 'No APs available to generate report'))
            return
        
        unique_floors = sorted(list({(ap['filename'], ap['floor']) for ap in self.aps_for_plotting}), key=self.floor_and_block_sort_key)
        
        order_dialog = tk.Toplevel(self.root)
        order_dialog.title(self.tr('Ordenar Pisos y Opciones de Informe', 'Sort Floors and Report Options'))
        order_dialog.geometry("800x600")

        main_frame = ttk.Frame(order_dialog, padding=10)
        main_frame.pack(fill=tk.BOTH, expand=True)

        left_frame = ttk.Frame(main_frame)
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5)

        right_frame = ttk.Frame(main_frame)
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=5)

        # Left side: Floor list and controls
        listbox_frame = ttk.Frame(left_frame)
        listbox_frame.pack(fill=tk.BOTH, expand=True)
        
        listbox = tk.Listbox(listbox_frame, selectmode=tk.SINGLE)
        scrollbar = ttk.Scrollbar(listbox_frame, orient="vertical", command=listbox.yview)
        listbox.configure(yscrollcommand=scrollbar.set)
        
        for filename, floor in unique_floors:
            listbox.insert(tk.END, f"{filename} - {self.tr('Piso', 'Floor')} {floor}")
        
        listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        def move_up():
            idx = listbox.curselection()
            if idx and idx[0] > 0:
                item = listbox.get(idx)
                listbox.delete(idx)
                listbox.insert(idx[0]-1, item)
                listbox.select_set(idx[0]-1)
        
        def move_down():
            idx = listbox.curselection()
            if idx and idx[0] < listbox.size()-1:
                item = listbox.get(idx)
                listbox.delete(idx)
                listbox.insert(idx[0]+1, item)
                listbox.select_set(idx[0]+1)
        
        btn_frame = ttk.Frame(left_frame)
        btn_frame.pack(pady=5)
        ttk.Button(btn_frame, text="↑ Subir", command=move_up).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="↓ Bajar", command=move_down).pack(side=tk.LEFT, padx=2)

        # Options
        options_frame = ttk.LabelFrame(left_frame, text="Opciones del Informe", padding=10)
        options_frame.pack(fill=tk.X, pady=10)

        # Cover Page Info
        ttk.Label(options_frame, text="Nombre del Cliente:").grid(row=0, column=0, sticky=tk.W, pady=2)
        self.client_name_entry = ttk.Entry(options_frame, width=30)
        self.client_name_entry.grid(row=0, column=1, sticky=tk.EW, pady=2)

        ttk.Label(options_frame, text="Ingeniero Responsable:").grid(row=1, column=0, sticky=tk.W, pady=2)
        self.engineer_name_entry = ttk.Entry(options_frame, width=30)
        self.engineer_name_entry.grid(row=1, column=1, sticky=tk.EW, pady=2)

        # Report Options
        self.plot_notes_var = tk.BooleanVar(value=True)
        self.include_note_totals_var = tk.BooleanVar(value=True)

        ttk.Checkbutton(options_frame, text="Graficar Notas", variable=self.plot_notes_var).grid(row=2, columnspan=2, sticky=tk.W, pady=(10, 0))
        ttk.Checkbutton(options_frame, text="Incluir Totales de Notas", variable=self.include_note_totals_var).grid(row=3, columnspan=2, sticky=tk.W)

        # Right side: Preview
        preview_frame = ttk.LabelFrame(right_frame, text=self.tr('Vista Previa del Plano', 'Floor Preview'), padding=10)
        preview_frame.pack(fill=tk.BOTH, expand=True)

        preview_canvas = tk.Canvas(preview_frame, bg='white')
        preview_canvas.pack(fill=tk.BOTH, expand=True)

        self.zoom_factor = 1.0
        self.pan_start_x = 0
        self.pan_start_y = 0
        self.pan_offset_x = 0
        self.pan_offset_y = 0
        self.canvas_img_id = None
        self.current_preview_floor = None
        self.preview_image_cache = None

        def zoom(event):
            # Determine the zoom point
            canvas_x, canvas_y = event.x, event.y
            
            # Convert canvas coordinates to image coordinates before zoom
            img_x_before = (canvas_x - self.pan_offset_x) / self.zoom_factor
            img_y_before = (canvas_y - self.pan_offset_y) / self.zoom_factor

            # Apply zoom
            if event.delta > 0:
                self.zoom_factor *= 1.1
            else:
                self.zoom_factor /= 1.1
            
            # Calculate image coordinates after zoom
            img_x_after = (canvas_x - self.pan_offset_x) / self.zoom_factor
            img_y_after = (canvas_y - self.pan_offset_y) / self.zoom_factor

            # Adjust pan offset to keep the zoom point stationary
            self.pan_offset_x += (img_x_after - img_x_before) * self.zoom_factor
            self.pan_offset_y += (img_y_after - img_y_before) * self.zoom_factor

            show_preview(None, update_notes=False)

        def start_pan(event):
            self.pan_start_x = event.x
            self.pan_start_y = event.y

        def pan(event):
            dx = event.x - self.pan_start_x
            dy = event.y - self.pan_start_y
            self.pan_offset_x += dx
            self.pan_offset_y += dy
            self.pan_start_x = event.x
            self.pan_start_y = event.y
            redraw_preview()

        preview_canvas.bind("<MouseWheel>", zoom)
        preview_canvas.bind("<ButtonPress-2>", start_pan)
        preview_canvas.bind("<B2-Motion>", pan)

        self.preview_notes = [] # This will hold notes for the currently displayed floor

        def redraw_preview():
            if not self.preview_image_cache:
                return

            preview_canvas.delete("all")
            
            # Resize the cached base image
            new_width = int(self.preview_image_cache.width * self.zoom_factor)
            new_height = int(self.preview_image_cache.height * self.zoom_factor)
            
            # Use a faster resize algorithm for interactive previews
            resized_img = self.preview_image_cache.resize((new_width, new_height), Image.LANCZOS)

            draw = ImageDraw.Draw(resized_img)
            
            # Draw notes from the current preview list
            for note in self.preview_notes:
                zoomed_note = note.copy()
                zoomed_note['x'] *= self.zoom_factor
                zoomed_note['y'] *= self.zoom_factor
                self.draw_note(draw, zoomed_note)

            self.tk_preview_img = ImageTk.PhotoImage(resized_img)
            
            self.canvas_img_id = preview_canvas.create_image(
                self.pan_offset_x, self.pan_offset_y, anchor=tk.NW, image=self.tk_preview_img
            )

        def show_preview(event, update_notes=True):
            selection = listbox.curselection()
            if not selection:
                return
            
            floor_entry = listbox.get(selection[0])
            filename, floor = self.parse_floor_entry(floor_entry)
            
            # If floor changes, reset view and load new data
            if (filename, floor) != self.current_preview_floor:
                self.current_preview_floor = (filename, floor)
                self.zoom_factor = 1.0
                self.pan_offset_x = 0
                self.pan_offset_y = 0
                
                # Load the base image for the new floor
                file_path = next((fp for fp in self.selected_files if os.path.basename(fp) == filename), None)
                if not file_path:
                    self.preview_image_cache = None
                    return
                
                temp_dir = tempfile.mkdtemp()
                img_path = self.generate_floor_image(file_path, floor, temp_dir, preview=True, plot_notes=False)
                if img_path:
                    # Open the image and immediately create a copy in memory
                    # to release the file lock before deleting the temp directory.
                    with Image.open(img_path) as img:
                        self.preview_image_cache = img.copy()
                else:
                    self.preview_image_cache = None
                
                # Now that the image is copied to memory, the temp dir can be safely deleted.
                shutil.rmtree(temp_dir)

                # Load notes for the new floor
                self.preview_notes = [
                    note.copy() for note in self.notes_data
                    if note['filename'] == filename and note['floor'] == floor
                ]

            redraw_preview()

        listbox.bind('<<ListboxSelect>>', show_preview)

        def confirm_order():
            self.floor_order = [listbox.get(i) for i in range(listbox.size())]
            self.plot_notes = self.plot_notes_var.get()
            self.include_note_totals = self.include_note_totals_var.get()
            self.client_name = self.client_name_entry.get()
            self.engineer_name = self.engineer_name_entry.get()
            order_dialog.destroy()
            self.create_word_document()
        
        ttk.Button(left_frame, text=self.tr('Generar Informe', 'Generate Report'), command=confirm_order).pack(pady=10)

    def create_word_document(self):
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=(("Word documents", "*.docx"), ("Todos los archivos", "*.*"))
        )
        if not save_path:
            return
        
        doc = Document()
        temp_dir = tempfile.mkdtemp()
        
        try:
            # --- Cover Page ---
            if os.path.exists(self.report_image_path):
                doc.add_picture(self.report_image_path, width=Inches(3.0))
            
            title = doc.add_paragraph()
            title.add_run('Informe de Cobertura Inalámbrica').bold = True
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph() # Spacer
            
            doc.add_paragraph(f"Cliente: {self.client_name}")
            doc.add_paragraph(f"Ingeniero Responsable: {self.engineer_name}")
            doc.add_paragraph(f"Fecha de Generación: {datetime.now().strftime('%Y-%m-%d')}")
            
            doc.add_page_break()

            # --- Data Aggregation ---
            summary_data = {}
            all_models = set()
            for filename, model, floor, count in self.ap_data:
                key = (filename, floor)
                if key not in summary_data:
                    summary_data[key] = {'aps': {}, 'notes': {"existente": 0, "nuevo": 0, "traslado": 0}}
                summary_data[key]['aps'][model] = summary_data[key]['aps'].get(model, 0) + count
                all_models.add(model)

            for filename, floor_counts in self.note_counts.items():
                for floor, counts in floor_counts.items():
                    key = (filename, floor)
                    if key not in summary_data:
                        summary_data[key] = {'aps': {}, 'notes': {"existente": 0, "nuevo": 0, "traslado": 0}}
                    summary_data[key]['notes'] = counts

            sorted_models = sorted(list(all_models))

            # --- Floor by Floor Section ---
            current_block = None
            for i, floor_entry in enumerate(self.floor_order):
                filename, floor = self.parse_floor_entry(floor_entry)
                
                if filename != current_block:
                    if current_block is not None:
                        doc.add_page_break()
                    current_block = filename
                    doc.add_heading(f"Bloque: {current_block}", level=1)

                file_path = next((fp for fp in self.selected_files if os.path.basename(fp) == filename), None)
                if not file_path: continue
                
                img_path = self.generate_floor_image(file_path, floor, temp_dir, preview=False)
                if not img_path: continue
                
                doc.add_heading(f"Piso: {floor}", level=2)
                doc.add_picture(img_path, width=Inches(6))
                
                # --- Per-Floor Summary Table ---
                doc.add_paragraph("\nResumen del Piso:", style='BodyText')
                floor_data = summary_data.get((filename, floor), {'aps': {}, 'notes': {}})
                ap_models_on_floor = floor_data.get('aps', {})
                note_counts_on_floor = floor_data.get('notes', {})

                if ap_models_on_floor or (self.include_note_totals and any(note_counts_on_floor.values())):
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Tipo'
                    hdr_cells[1].text = 'Detalle'
                    hdr_cells[2].text = 'Cantidad'
                    hdr_cells[3].text = 'Total'

                    # AP Rows
                    ap_models_sorted = sorted(ap_models_on_floor.items())
                    if ap_models_sorted:
                        total_aps_floor = sum(ap_models_on_floor.values())
                        start_ap_row_idx = len(table.rows)
                        for i, (model, count) in enumerate(ap_models_sorted):
                            row_cells = table.add_row().cells
                            if i == 0:
                                row_cells[0].text = 'APs'
                                row_cells[3].text = str(total_aps_floor)
                            row_cells[1].text = model
                            row_cells[2].text = str(count)
                        
                        if len(ap_models_sorted) > 1:
                            table.cell(start_ap_row_idx, 0).merge(table.cell(len(table.rows) - 1, 0))
                            table.cell(start_ap_row_idx, 3).merge(table.cell(len(table.rows) - 1, 3))

                    # Note Rows
                    if self.include_note_totals:
                        note_types_sorted = ["existente", "nuevo", "traslado"]
                        total_notes_floor = sum(note_counts_on_floor.get(nt, 0) for nt in note_types_sorted)
                        
                        if total_notes_floor > 0 or not ap_models_sorted:
                            start_note_row_idx = len(table.rows)
                            for i, note_type in enumerate(note_types_sorted):
                                count = note_counts_on_floor.get(note_type, 0)
                                row_cells = table.add_row().cells
                                if i == 0:
                                    row_cells[0].text = 'Notas'
                                    row_cells[3].text = str(total_notes_floor)
                                row_cells[1].text = note_type.capitalize()
                                row_cells[2].text = str(count)
                            
                            if len(note_types_sorted) > 1:
                                table.cell(start_note_row_idx, 0).merge(table.cell(len(table.rows) - 1, 0))
                                table.cell(start_note_row_idx, 3).merge(table.cell(len(table.rows) - 1, 3))

                if i < len(self.floor_order) - 1:
                    # Check if the next floor is in a different block
                    next_filename, _ = self.parse_floor_entry(self.floor_order[i+1])
                    if filename != next_filename:
                        # Don't add a page break if it's the last floor of a block before the summary
                        pass
                    else:
                        doc.add_page_break()

            # --- General Summary Section ---
            doc.add_page_break()
            doc.add_heading("Resumen General del Proyecto", level=1)
            
            cols = 2 + len(sorted_models) + (3 if self.include_note_totals else 0)
            table = doc.add_table(rows=1, cols=cols)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Archivo'
            hdr_cells[1].text = 'Piso'
            for i, model in enumerate(sorted_models):
                hdr_cells[2 + i].text = model
            if self.include_note_totals:
                hdr_cells[2 + len(sorted_models)].text = 'Existente'
                hdr_cells[3 + len(sorted_models)].text = 'Nuevo'
                hdr_cells[4 + len(sorted_models)].text = 'Traslado'

            total_aps_by_model = {model: 0 for model in sorted_models}
            total_notes = {"existente": 0, "nuevo": 0, "traslado": 0}

            for (filename, floor), data in sorted(summary_data.items(), key=lambda item: (item[0][0], self.floor_sort_key(item[0][1]))):
                row_cells = table.add_row().cells
                row_cells[0].text = filename
                row_cells[1].text = str(floor)
                for i, model in enumerate(sorted_models):
                    count = data['aps'].get(model, 0)
                    row_cells[2 + i].text = str(count)
                    total_aps_by_model[model] += count
                if self.include_note_totals:
                    notes = data['notes']
                    row_cells[2 + len(sorted_models)].text = str(notes.get('existente', 0))
                    row_cells[3 + len(sorted_models)].text = str(notes.get('nuevo', 0))
                    row_cells[4 + len(sorted_models)].text = str(notes.get('traslado', 0))
                    total_notes['existente'] += notes.get('existente', 0)
                    total_notes['nuevo'] += notes.get('nuevo', 0)
                    total_notes['traslado'] += notes.get('traslado', 0)

            # Total Row
            total_row = table.add_row().cells
            total_row[0].text = 'TOTAL GENERAL'
            total_row[0].merge(total_row[1])
            for i, model in enumerate(sorted_models):
                total_row[2 + i].text = str(total_aps_by_model[model])
            if self.include_note_totals:
                total_row[2 + len(sorted_models)].text = str(total_notes['existente'])
                total_row[3 + len(sorted_models)].text = str(total_notes['nuevo'])
                total_row[4 + len(sorted_models)].text = str(total_notes['traslado'])

            # --- Charts Section ---
            doc.add_page_break()
            doc.add_heading("Visualización de Datos", level=1)

            building_summary = {}
            for (filename, _), data in summary_data.items():
                if filename not in building_summary:
                    building_summary[filename] = {'models': {m: 0 for m in sorted_models}, 'notes': {'existente': 0, 'nuevo': 0, 'traslado': 0}}
                for model, count in data['aps'].items():
                    building_summary[filename]['models'][model] += count
                for note_type, count in data['notes'].items():
                    building_summary[filename]['notes'][note_type] += count
            
            # Chart 1: AP Models per Building
            doc.add_heading("Total de Modelos de AP por Edificio", level=2)
            model_chart_path = self.generate_bar_chart(
                data={b: d['models'] for b, d in building_summary.items()},
                title="Cantidad de APs por Modelo y Edificio",
                xlabel="Edificio", ylabel="Cantidad de APs",
                output_dir=temp_dir, filename="model_chart.png"
            )
            if model_chart_path: doc.add_picture(model_chart_path, width=Inches(6))

            # Chart 2: Notes per Building
            if self.include_note_totals:
                doc.add_heading("Total de Notas por Edificio", level=2)
                note_chart_path = self.generate_bar_chart(
                    data={b: d['notes'] for b, d in building_summary.items()},
                    title="Cantidad de Notas por Tipo y Edificio",
                    xlabel="Edificio", ylabel="Cantidad de Notas",
                    output_dir=temp_dir, filename="note_chart.png"
                )
                if note_chart_path: doc.add_picture(note_chart_path, width=Inches(6))

            doc.save(save_path)
            messagebox.showinfo("Éxito", "Informe generado correctamente")
        except Exception as e:
            messagebox.showerror("Error", f"Error al generar informe:\n{str(e)}")
        finally:
            shutil.rmtree(temp_dir)
    def generate_bar_chart(self, data, title, xlabel, ylabel, output_dir, filename):
        try:
            buildings = list(data.keys())
            if not buildings:
                return None

            categories = sorted(list(next(iter(data.values())).keys()))
            if not categories:
                return None

            num_buildings = len(buildings)
            num_categories = len(categories)

            # Dynamic sizing
            fig_width = max(10, num_buildings * num_categories * 0.2)
            fig_height = 6
            
            fig, ax = plt.subplots(figsize=(fig_width, fig_height))
            
            x = np.arange(num_buildings)
            width = 0.8 / num_categories
            
            # Use a color palette
            colors = plt.cm.get_cmap('viridis', num_categories)

            for i, category in enumerate(categories):
                counts = [d.get(category, 0) for d in data.values()]
                offset = width * (i - (num_categories - 1) / 2)
                rects = ax.bar(x + offset, counts, width, label=category, color=colors(i))
                ax.bar_label(rects, padding=3, fontsize=8)

            ax.set_ylabel(ylabel, fontsize=10)
            ax.set_xlabel(xlabel, fontsize=10)
            ax.set_title(title, fontsize=12, weight='bold')
            ax.set_xticks(x)
            ax.set_xticklabels(buildings, rotation=45, ha="right", fontsize=9)
            
            # Adjust legend position
            ax.legend(title='Categorías', bbox_to_anchor=(1.04, 1), loc='upper left')
            
            # Add grid for better readability
            ax.yaxis.grid(True, linestyle='--', alpha=0.7)

            fig.tight_layout(rect=[0, 0, 0.85, 1]) # Adjust layout to make room for legend
            
            output_path = os.path.join(output_dir, filename)
            plt.savefig(output_path, bbox_inches='tight')
            plt.close(fig)
            return output_path
        except Exception as e:
            print(f"Error generating chart {filename}: {e}")
            return None

    def parse_floor_entry(self, entry):
        parts = entry.rsplit(" - ", 1)
        if len(parts) != 2:
            return entry.strip(), ""

        floor_part = parts[1].strip()
        for token in ("Piso ", "Floor "):
            if floor_part.startswith(token):
                floor_part = floor_part[len(token):]
                break

        return parts[0].strip(), floor_part.strip()

    def floor_and_block_sort_key(self, floor_entry):
        filename, floor_str = floor_entry
        
        # Extract block number from filename
        block_match = re.search(r'bloque_(\d+)', filename, re.IGNORECASE)
        block_number = int(block_match.group(1)) if block_match else 0
        
        # Extract floor number
        floor_number = self.floor_sort_key(floor_str)
        
        return (block_number, floor_number)

    def generate_floor_image(self, file_path, floor_number, output_dir, preview=False, plot_notes=True):
        try:
            filename = os.path.basename(file_path)  # Obtener el nombre del archivo desde la ruta
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                floor_plans = self.load_json_from_zip(zip_ref, 'floorPlans.json').get('floorPlans', [])
                
                for floor in floor_plans:
                    if self.extract_floor_number(floor.get('name', '')) == floor_number:
                        image_id = floor.get('imageId')
                        if not image_id:
                            return None
                        
                        with zip_ref.open(f"image-{image_id}") as img_file:
                            image = Image.open(img_file).convert("RGB")
                        
                        draw = ImageDraw.Draw(image)
                        aps = [ap for ap in self.aps_for_plotting 
                            if ap['filename'] == filename  # Usar el nombre del archivo
                            and ap['floor'] == floor_number]
                        
                        for ap in aps:
                            x = ap['x']
                            y = ap['y']
                            
                            draw.ellipse((x-self.circle_radius, y-self.circle_radius,
                                        x+self.circle_radius, y+self.circle_radius),
                                    fill=self.circle_color, outline="black")
                            
                            try:
                                font = ImageFont.truetype("arialbd.ttf", self.font_size)
                            except IOError:
                                font = ImageFont.load_default()
                            
                            text_position = (x + self.circle_radius + 2, y - self.circle_radius)
                            
                            if self.outline_width > 0:
                                for dx in range(-self.outline_width, self.outline_width+1):
                                    for dy in range(-self.outline_width, self.outline_width+1):
                                        if dx != 0 or dy != 0:
                                            draw.text(
                                                (text_position[0] + dx, text_position[1] + dy),
                                                ap['name'],
                                                fill=self.outline_color,
                                                font=font
                                            )
                            
                            draw.text(text_position, ap['name'], 
                                    fill=self.text_color, font=font)
                        
                        if plot_notes and ((not preview and self.plot_notes) or (preview and self.plot_notes_var.get())):
                            # Get AP label bounding boxes to avoid overlap
                            ap_bboxes = []
                            font = ImageFont.truetype("arialbd.ttf", self.font_size)
                            for ap in aps:
                                text_position = (ap['x'] + self.circle_radius + 2, ap['y'] - self.circle_radius)
                                # Use textbbox for accurate measurement
                                bbox = draw.textbbox(text_position, ap['name'], font=font)
                                ap_bboxes.append(bbox)

                            floor_notes = [n for n in self.notes_data if n['filename'] == filename and n['floor'] == floor_number]
                            
                            for note in floor_notes:
                                note_font = ImageFont.truetype("arial.ttf", self.note_font_size)
                                note_bbox = draw.textbbox((note['x'], note['y']), note['text'], font=note_font)

                                # Collision detection and avoidance
                                attempts = 0
                                while any(self.bboxes_overlap(note_bbox, ap_bbox) for ap_bbox in ap_bboxes) and attempts < 20:
                                    # Move note down and to the right to find a clear spot
                                    note['x'] += 5
                                    note['y'] += 5
                                    note_bbox = draw.textbbox((note['x'], note['y']), note['text'], font=note_font)
                                    attempts += 1
                                
                                self.draw_note(draw, note)
                        
                        safe_name = re.sub(r'[\\/*?:"<>|]', '_', floor.get('name', ''))
                        output_path = os.path.join(output_dir, f"{filename}_{safe_name}.png")
                        image.save(output_path)
                        return output_path
            return None
        except Exception as e:
            print(f"Error generando imagen: {str(e)}")
            return None

    def bboxes_overlap(self, bbox1, bbox2):
        # Checks if two bounding boxes (x0, y0, x1, y1) overlap
        return not (bbox1[2] < bbox2[0] or bbox1[0] > bbox2[2] or
                    bbox1[3] < bbox2[1] or bbox1[1] > bbox2[3])
            

    def merge_existing_pdfs(self):
        file_paths = filedialog.askopenfilenames(
            title=self.tr('Seleccionar archivos PDF para unir', 'Select PDF files to merge'),
            filetypes=(("PDF files", "*.pdf"), ("Todos los archivos", "*.*"))
        )

        if not file_paths:
            return

        save_path = filedialog.asksaveasfilename(
            title=self.tr('Guardar PDF combinado como...','Save merged PDF as...'),
            defaultextension=".pdf",
            filetypes=(("PDF files", "*.pdf"), ("Todos los archivos", "*.*"))
        )

        if not save_path:
            return

        try:
            pdf_merger = PdfWriter()
            for path in file_paths:
                pdf_merger.append(path)

            with open(save_path, "wb") as f_out:
                pdf_merger.write(f_out)
            
            pdf_merger.close()
            messagebox.showinfo("Éxito", f"PDF combinado guardado en:\n{save_path}")

        except Exception as e:
            messagebox.showerror("Error", f"Ocurrió un error al unir los PDFs:\n{str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = WirelessSurveyExtractor(root)
    root.mainloop()







